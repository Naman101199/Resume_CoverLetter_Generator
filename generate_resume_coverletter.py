#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Jan 07 01:51:35 2023

@author: naman
"""

resume_coverletter_directory  = '/Users/naman/Library/CloudStorage/OneDrive-TheUniversityofChicago/Resume/'

resume_original_file_name = 'Naman Mehta Resume'
coverletter_original_file_name = 'Naman Mehta Cover Letter'

company_name = 'test'
position_name = 'Data Science Intern'

role_types = ['Finance','Data Science']
cover_letter_by_type = ['Naman Mehta Cover Letter Quant.docx', 'Naman Mehta Cover Letter Data Science.docx']
applying_to_role_index = 1


from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import shutil
import os

os.chdir(resume_coverletter_directory)
if company_name not in os.listdir():
    os.mkdir(company_name)

fileNameResume = company_name + '/' + resume_original_file_name + ' ' + company_name + ' ' + position_name +'.pdf'
shutil.copy(resume_original_file_name + '.pdf', fileNameResume)

print(resume_original_file_name + '.pdf' , fileNameResume)

path_cover_letter = ''

if applying_to_role_index == 0:
    document = Document(cover_letter_by_type[0])
else:
    document = Document(cover_letter_by_type[1])
# Create the PDF document

fileNameCoverLetter = company_name + "/" + coverletter_original_file_name + ' ' + company_name + ' ' + position_name + ".pdf"
pdf = SimpleDocTemplate(fileNameCoverLetter, pagesize=letter)
pdf_elements = []

# Iterate through the paragraphs in the Word document
for p in document.paragraphs:
    if 'company_name' in p.text:
        p.text = p.text.replace('company_name', company_name)
    if 'position_name' in p.text:
        p.text = p.text.replace('position_name', position_name)
    # Get the paragraph style information
    style = p.style
    font = p.style.font
    alignment = p.alignment
    # Create a Paragraph object with the style information
    pdf_elements.append(Paragraph(p.text, 
                                 style=ParagraphStyle(name='Times-Roman',
                                                      fontName = 'Times-Roman',
                                                      fontSize = 12)))
    pdf_elements.append(Spacer(1,12))

pdf.build(pdf_elements)